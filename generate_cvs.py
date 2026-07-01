import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_en_cv():
    doc = docx.Document()
    
    name = doc.add_heading('HASSAN MUNDHER ABDULWAHID', 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph('Basra, Iraq | +964 783 895 3531 | hasanmthr02@gmail.com\nPortfolio: civil-eng-pro.vercel.app | LinkedIn: linkedin.com/in/hasanmundher')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('PROFESSIONAL PROFILE', level=1)
    doc.add_paragraph('Dedicated Civil Engineer with a strong foundation in structural design and site execution. Proficient in modern engineering software including STAAD.Pro, ETABS, and Revit. Experienced in overseeing residential construction phases and ensuring quality control. Currently managing administrative and client-facing responsibilities as a Contract Sales Employee at the Ministry of Electricity, demonstrating excellent communication, adaptability, and organizational skills. Committed to continuous professional development and strict adherence to safety standards.')
    
    doc.add_heading('CORE COMPETENCIES', level=1)
    doc.add_paragraph('Structural Engineering: STAAD.Pro, ETABS, SAP2000, SAFE', style='List Bullet')
    doc.add_paragraph('BIM & Drafting: AutoCAD (2D/3D), Revit Architecture, Revit Structure', style='List Bullet')
    doc.add_paragraph('Construction Management: Site Surveying (Auto Level), Quantity Surveying, QC/QA', style='List Bullet')
    doc.add_paragraph('Soft Skills: Problem Solving, Customer Service, Team Collaboration, MS Office', style='List Bullet')
    
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Ministry of Electricity').bold = True
    p.add_run(' | Basra, Iraq\n')
    p.add_run('Sales Representative (Contract)').italic = True
    p.add_run('\n- Handle daily operational sales tasks and ensure efficient service delivery to citizens.\n- Organize and review revenue data, maintaining accurate financial records.\n- Resolve customer complaints and streamline administrative workflows.')
    
    p = doc.add_paragraph()
    p.add_run('Al-Farah Residential Complex').bold = True
    p.add_run(' | Basra, Iraq\n')
    p.add_run('Site Engineer (Intern)').italic = True
    p.add_run(' | July 2025\n- Conducted field supervision of construction phases to ensure compliance with approved engineering blueprints.\n- Monitored interior finishing works and inspected the quality of materials used on-site.\n- Collaborated with technical teams to quickly resolve on-site engineering challenges.')
    
    p = doc.add_paragraph()
    p.add_run('Private Residential Projects').bold = True
    p.add_run(' | Basra, Iraq\n')
    p.add_run('Site Engineer (Trainee)').italic = True
    p.add_run(' | 2023 - 2025\n- Utilized Auto Level instruments accurately to establish elevations during foundation casting.\n- Executed quantity takeoffs for construction materials and documented daily site progress.\n- Enforced occupational safety guidelines to ensure a secure working environment.')
    
    doc.add_heading('EDUCATION', level=1)
    p = doc.add_paragraph()
    p.add_run('Bachelor of Science in Civil Engineering').bold = True
    p.add_run(' | Expected 2026\nShatt Al-Arab University, College of Engineering | Basra, Iraq')
    
    doc.add_heading('TECHNICAL PROJECTS', level=1)
    p = doc.add_paragraph()
    p.add_run('Graduation Project: Multi-Story Steel Structure Design').bold = True
    p.add_run('\n- Designed a complete steel structure framework using STAAD.Pro.\n- Ensured compliance with AISC codes and produced detailed architectural and structural shop drawings.')
    
    p = doc.add_paragraph()
    p.add_run('CivilEng - Mobile Application').bold = True
    p.add_run('\n- Developed a customized Android application designed to assist field engineers in automating quantity estimations for concrete, steel, and finishing materials.')
    
    doc.add_heading('CERTIFICATIONS', level=1)
    doc.add_paragraph('Occupational Safety and Health (OSHA) - Microtech', style='List Bullet')
    doc.add_paragraph('Civil Quality Control & Assurance (QC/QA) - Microtech', style='List Bullet')
    doc.add_paragraph('ISO 9001:2015 Non-conformances and Corrective Actions - Alison', style='List Bullet')
    doc.add_paragraph('Diploma in MS Project for Civil Engineers - Alison', style='List Bullet')
    
    doc.save('Hassan_Mundher_CV_EN.docx')

def create_ar_cv():
    doc = docx.Document()
    
    name = doc.add_heading('حسن منذر عبدالواحد', 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph('البصرة، العراق | +964 783 895 3531 | hasanmthr02@gmail.com\nمعرض الأعمال: civil-eng-pro.vercel.app | لينكد إن: linkedin.com/in/hasanmundher')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('النبذة المهنية', level=1)
    doc.add_paragraph('مهندس مدني طموح، أمتلك أساساً متيناً في التصميم الإنشائي والإشراف الميداني. أجيد استخدام برامج الهندسة الحديثة مثل STAAD.Pro و ETABS و Revit. لدي خبرة عملية في متابعة مراحل التشييد للمشاريع السكنية وضمان جودة التنفيذ. أعمل حالياً كموظف عقود في وزارة الكهرباء (قسم المبيعات)، مما عزز من مهاراتي في التواصل الإداري، التنظيم، وحل المشكلات. ملتزم بالتطوير المهني المستمر وتطبيق أعلى معايير السلامة والجودة.')
    
    doc.add_heading('المهارات الأساسية', level=1)
    doc.add_paragraph('الهندسة الإنشائية: إجادة برامج STAAD.Pro, ETABS, SAP2000, SAFE', style='List Bullet')
    doc.add_paragraph('الرسم والنمذجة: احتراف العمل على AutoCAD و Revit Architecture و Revit Structure', style='List Bullet')
    doc.add_paragraph('إدارة المواقع: المسح بجهاز (Auto Level)، حصر الكميات، ومراقبة الجودة (QC/QA)', style='List Bullet')
    doc.add_paragraph('المهارات الإدارية: إدارة المبيعات، خدمة العملاء، حل المشكلات، واستخدام برامج MS Office', style='List Bullet')
    
    doc.add_heading('الخبرات العملية', level=1)
    
    p = doc.add_paragraph()
    p.add_run('وزارة الكهرباء').bold = True
    p.add_run(' | البصرة، العراق\n')
    p.add_run('موظف عقود (قسم المبيعات)').italic = True
    p.add_run('\n- الإشراف على إنجاز المعاملات اليومية ومتابعة تقديم الخدمات للمشتركين بكفاءة عالية.\n- تنظيم بيانات المبيعات ومراجعة الإيرادات بشكل دوري لضمان دقة السجلات المالية.\n- معالجة طلبات وشكاوى المواطنين وضمان سير العمل الإداري بسلاسة.')
    
    p = doc.add_paragraph()
    p.add_run('مجمع الفرح السكني').bold = True
    p.add_run(' | البصرة، العراق\n')
    p.add_run('مهندس موقع (متدرب)').italic = True
    p.add_run(' | يوليو 2025\n- المتابعة الميدانية لمراحل التشييد والبناء لضمان مطابقة العمل للمخططات الهندسية المعتمدة.\n- الإشراف على أعمال التشطيبات الداخلية والتأكد من جودة المواد المستخدمة في الموقع.\n- التعاون المباشر مع الكادر الفني والهندسي لحل أي مشكلات طارئة أثناء التنفيذ.')
    
    p = doc.add_paragraph()
    p.add_run('مشاريع سكنية خاصة').bold = True
    p.add_run(' | البصرة، العراق\n')
    p.add_run('مهندس موقع (فترة تدريبية)').italic = True
    p.add_run(' | 2023 - 2025\n- استخدام أجهزة المساحة (Auto Level) لضبط المناسيب بدقة في مراحل صب الأساسات.\n- حصر الكميات للمواد الإنشائية وتوثيق نسب إنجاز العمل اليومية في سجلات الموقع.\n- تطبيق إرشادات السلامة المهنية لحماية العاملين وضمان بيئة عمل آمنة.')
    
    doc.add_heading('التعليم الأكاديمي', level=1)
    p = doc.add_paragraph()
    p.add_run('بكالوريوس في الهندسة المدنية').bold = True
    p.add_run(' | سنة التخرج المتوقعة: 2026\nجامعة شط العرب، كلية الهندسة | البصرة، العراق')
    
    doc.add_heading('المشاريع التقنية', level=1)
    p = doc.add_paragraph()
    p.add_run('مشروع التخرج: تصميم مبنى فولاذي متعدد الطوابق').bold = True
    p.add_run('\n- إعداد تصميم إنشائي متكامل لهيكل فولاذي باستخدام برنامج STAAD.Pro.\n- التأكد من مطابقة التصميم لمعايير كود AISC، وإنتاج المخططات التنفيذية التفصيلية.')
    
    p = doc.add_paragraph()
    p.add_run('تطبيق CivilEng للهواتف الذكية').bold = True
    p.add_run('\n- برمجة وتطوير تطبيق أندرويد موجه للمهندسين الميدانيين لتسهيل وتسريع حساب كميات الخرسانة، الحديد، ومواد التشطيبات بصورة دقيقة.')
    
    doc.add_heading('الشهادات والدورات المعتمدة', level=1)
    doc.add_paragraph('دورة في الصحة والسلامة المهنية (OSHA) - شركة مايكرو تكنك', style='List Bullet')
    doc.add_paragraph('دورة في مراقبة وضمان الجودة المدنية (QC/QA) - شركة مايكرو تكنك', style='List Bullet')
    doc.add_paragraph('شهادة في الإجراءات التصحيحية لنظام إدارة الجودة ISO 9001 - منصة Alison', style='List Bullet')
    doc.add_paragraph('دبلوم في استخدام برنامج MS Project للمهندسين - منصة Alison', style='List Bullet')
    
    # Try to set right-to-left orientation for the whole document paragraphs if possible
    # We will just save it as is; Arabic text automatically displays correctly in word when language is Arabic.
    doc.save('Hassan_Mundher_CV_AR.docx')

create_en_cv()
create_ar_cv()
